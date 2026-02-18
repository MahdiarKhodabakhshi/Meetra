from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import BinaryIO
from urllib.parse import urlparse

from celery.utils.log import get_task_logger
from sqlalchemy.orm import Session

from app.db import SessionLocal
from app.models import IngestionJob, Profile, ResumeVersion, User
from app.models.resume_version import ResumeVersionStatus
from app.models.user import UserStatus
from app.storage.factory import get_storage
from app.worker.celery_app import celery_app

logger = get_task_logger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

HEADING_ALIASES = {
    "summary": {"summary", "professional summary", "profile", "about"},
    "skills": {"skills", "technical skills", "core skills", "competencies", "tech stack"},
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment history",
    },
    "education": {"education", "academic background", "qualifications"},
    "industries": {"industries", "industry", "industry exposure"},
}
TITLE_KEYWORDS = (
    "engineer",
    "developer",
    "manager",
    "analyst",
    "designer",
    "consultant",
    "director",
    "lead",
    "architect",
    "scientist",
    "specialist",
    "coordinator",
    "administrator",
    "product",
    "owner",
)
KNOWN_SKILLS = (
    "Python",
    "Java",
    "JavaScript",
    "TypeScript",
    "SQL",
    "PostgreSQL",
    "MySQL",
    "AWS",
    "Azure",
    "Docker",
    "Kubernetes",
    "React",
    "Node.js",
    "FastAPI",
    "Django",
    "Flask",
    "Git",
    "Linux",
    "Pandas",
    "NumPy",
    "Machine Learning",
    "Data Analysis",
    "REST APIs",
    "GraphQL",
    "Microservices",
    "Agile",
    "Scrum",
    "Project Management",
)
INDUSTRY_KEYWORDS = {
    "FinTech": ("fintech", "bank", "banking", "payments", "insurance"),
    "Healthcare": ("healthcare", "hospital", "clinical", "medical", "ehr"),
    "Education": ("education", "edtech", "university", "school", "learning"),
    "E-commerce": ("e-commerce", "ecommerce", "retail", "marketplace"),
    "SaaS": ("saas", "subscription software", "b2b software"),
    "Government": ("government", "public sector", "ministry"),
    "Telecom": ("telecom", "telecommunications", "network operations"),
    "Energy": ("energy", "oil", "gas", "renewable", "utilities"),
}
SPLIT_RE = re.compile(r"[,\n;|]+")
WHITESPACE_RE = re.compile(r"\s+")
HEADING_PUNCT_RE = re.compile(r"[:\- ]+$")
NON_HEADING_CHARS_RE = re.compile(r"[^a-z ]")

USER_CONFIRMED = "USER_CONFIRMED"
MANUAL_OVERRIDES_KEY = "manual_overrides"
MANUALLY_EDITABLE_PROFILE_FIELDS = {
    "headline",
    "summary",
    "skills",
    "titles",
    "industries",
}


@celery_app.task(name="ingest_resume_text")
def ingest_resume_text(user_id: str | None, raw_text: str) -> dict:
    # Create DB row
    db: Session = SessionLocal()
    try:
        job = IngestionJob(
            user_id=uuid.UUID(user_id) if user_id else None,
            kind="resume_text",
            status="processing",
            input_text=raw_text,
        )
        db.add(job)
        db.commit()
        db.refresh(job)

        logger.info("ingest_resume_text started job_id=%s user_id=%s", job.id, user_id)

        # Placeholder work (no PDF parsing yet)
        # Later: parse, embed, match, etc.

        job.status = "completed"
        db.commit()

        logger.info("ingest_resume_text completed job_id=%s", job.id)
        return {"job_id": str(job.id), "status": job.status}
    except Exception as e:
        db.rollback()
        # Best-effort failure record (if job exists)
        try:
            if "job" in locals():
                job.status = "failed"
                job.error = str(e)
                db.commit()
        except Exception:
            pass
        raise
    finally:
        db.close()


@dataclass(slots=True)
class StructuredResume:
    headline: str | None
    summary: str | None
    skills: list[str]
    titles: list[str]
    industries: list[str]
    education_json: dict
    experience_json: dict
    keywords: list[str]
    confidence_json: dict[str, float]
    parse_confidence: float


class ResumeTaskError(Exception):
    def __init__(self, code: str, message: str) -> None:
        self.code = code
        self.message = message
        super().__init__(message)


def _manual_override_fields(profile: Profile | None) -> set[str]:
    if profile is None or not isinstance(profile.confidence_json, dict):
        return set()

    confidence = profile.confidence_json
    overrides: set[str] = set()
    raw_overrides = confidence.get(MANUAL_OVERRIDES_KEY, [])
    if isinstance(raw_overrides, list):
        for field in raw_overrides:
            field_name = str(field)
            if field_name in MANUALLY_EDITABLE_PROFILE_FIELDS:
                overrides.add(field_name)

    for field in MANUALLY_EDITABLE_PROFILE_FIELDS:
        field_confidence = confidence.get(field)
        if isinstance(field_confidence, dict):
            source = str(field_confidence.get("source", "")).upper()
            if source == USER_CONFIRMED:
                overrides.add(field)

    return overrides


def _merge_confidence_json(
    parsed_confidence: dict[str, float],
    existing_confidence: dict | None,
    manual_fields: set[str],
) -> dict:
    merged = dict(parsed_confidence)
    if isinstance(existing_confidence, dict):
        for key, value in existing_confidence.items():
            if key in MANUALLY_EDITABLE_PROFILE_FIELDS or key == MANUAL_OVERRIDES_KEY:
                continue
            merged[key] = value

    if manual_fields:
        merged[MANUAL_OVERRIDES_KEY] = sorted(manual_fields)
        for field in manual_fields:
            merged[field] = {"value": 1.0, "source": USER_CONFIRMED}
    return merged


def _storage_key_from_uri(uri: str) -> str:
    parsed = urlparse(uri)
    if not parsed.scheme:
        return uri.lstrip("/")
    key = f"{parsed.netloc}{parsed.path}"
    return key.lstrip("/")


def _looks_like_pdf(resume: ResumeVersion) -> bool:
    return resume.mime_type.lower() == PDF_MIME or (
        Path(resume.original_filename).suffix.lower() == ".pdf"
    )


def _looks_like_docx(resume: ResumeVersion) -> bool:
    return resume.mime_type.lower() == DOCX_MIME or (
        Path(resume.original_filename).suffix.lower() == ".docx"
    )


def _extract_pdf_text(file_obj: BinaryIO) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ResumeTaskError("TEXT_EXTRACTION_FAILED", "pdfplumber is not installed") from exc

    try:
        with pdfplumber.open(file_obj) as pdf:
            pages: list[str] = []
            for page in pdf.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)
    except Exception as exc:
        raise ResumeTaskError("TEXT_EXTRACTION_FAILED", f"failed to parse pdf: {exc}") from exc


def _extract_docx_text(file_obj: BinaryIO) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ResumeTaskError("TEXT_EXTRACTION_FAILED", "python-docx is not installed") from exc

    try:
        doc = docx.Document(file_obj)
        chunks: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    chunks.append(" | ".join(values))
        return "\n".join(chunks)
    except Exception as exc:
        raise ResumeTaskError("TEXT_EXTRACTION_FAILED", f"failed to parse docx: {exc}") from exc


def _extract_resume_text(resume: ResumeVersion) -> str:
    storage = get_storage()
    key = _storage_key_from_uri(resume.file_uri)
    try:
        with storage.open(key) as file_obj:
            if _looks_like_pdf(resume):
                text = _extract_pdf_text(file_obj)
            elif _looks_like_docx(resume):
                text = _extract_docx_text(file_obj)
            else:
                raise ResumeTaskError(
                    "TEXT_EXTRACTION_FAILED",
                    f"unsupported file type: {resume.mime_type}",
                )
    except ResumeTaskError:
        raise
    except FileNotFoundError as exc:
        raise ResumeTaskError(
            "TEXT_EXTRACTION_FAILED",
            "uploaded file is missing in storage",
        ) from exc
    except Exception as exc:
        raise ResumeTaskError(
            "TEXT_EXTRACTION_FAILED",
            f"unable to open resume file: {exc}",
        ) from exc

    if not text.strip():
        raise ResumeTaskError("TEXT_EXTRACTION_FAILED", "no text extracted from resume")
    return text


def _normalize_line(value: str) -> str:
    return WHITESPACE_RE.sub(" ", value.strip())


def _dedupe_keep_order(items: list[str], *, key_norm=str.lower) -> list[str]:
    output: list[str] = []
    seen: set[str] = set()
    for item in items:
        normalized = key_norm(item)
        if normalized in seen:
            continue
        output.append(item)
        seen.add(normalized)
    return output


def _heading_name(line: str) -> str | None:
    candidate = line.strip().lower()
    if not candidate:
        return None
    candidate = HEADING_PUNCT_RE.sub("", candidate)
    candidate = NON_HEADING_CHARS_RE.sub(" ", candidate)
    candidate = _normalize_line(candidate)
    if not candidate or len(candidate.split()) > 4:
        return None
    for section, aliases in HEADING_ALIASES.items():
        if candidate in aliases:
            return section
    return None


def _split_sections(text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {"preamble": []}
    current = "preamble"
    for raw_line in text.replace("\r\n", "\n").splitlines():
        stripped = raw_line.strip()
        if stripped:
            heading = _heading_name(stripped)
            if heading is not None:
                current = heading
                sections.setdefault(current, [])
                continue
        sections.setdefault(current, []).append(raw_line)

    compact: dict[str, str] = {}
    for key, lines in sections.items():
        normalized = "\n".join(_normalize_line(line) for line in lines if line.strip())
        compact[key] = normalized.strip()
    return compact


def _section_to_items(section_text: str, *, max_items: int = 20) -> list[str]:
    if not section_text:
        return []
    raw_items = [_normalize_line(token) for token in SPLIT_RE.split(section_text)]
    items: list[str] = []
    for item in raw_items:
        if not item:
            continue
        if len(item) > 120:
            continue
        if item.count(" ") > 12:
            continue
        items.append(item.strip(" -\u2022"))
    return _dedupe_keep_order(items)[:max_items]


def _extract_skills(sections: dict[str, str], full_text: str) -> tuple[list[str], float]:
    from_section = _section_to_items(sections.get("skills", ""), max_items=40)
    if from_section:
        return from_section, 0.85

    lower_text = full_text.lower()
    inferred = [skill for skill in KNOWN_SKILLS if skill.lower() in lower_text]
    inferred = _dedupe_keep_order(inferred)
    if inferred:
        return inferred[:40], 0.55
    return [], 0.3


def _looks_like_contact_line(line: str) -> bool:
    lower = line.lower()
    return (
        "@" in line
        or "linkedin.com" in lower
        or "github.com" in lower
        or re.search(r"\+?\d[\d\- ()]{6,}", line) is not None
    )


def _extract_titles(
    sections: dict[str, str],
    full_text: str,
) -> tuple[list[str], float]:
    source = sections.get("experience", "") or sections.get("preamble", "")
    lines = [line for line in source.splitlines() if line.strip()]
    titles: list[str] = []
    for line in lines:
        normalized = _normalize_line(line)
        if _looks_like_contact_line(normalized):
            continue
        lower = normalized.lower()
        if not any(keyword in lower for keyword in TITLE_KEYWORDS):
            continue
        candidate = re.split(r"\s+at\s+|\s+\|\s+|\s+-\s+", normalized, maxsplit=1)[0].strip()
        if 2 <= len(candidate) <= 90:
            titles.append(candidate)

    titles = _dedupe_keep_order(titles)[:15]
    if titles:
        confidence = 0.75 if sections.get("experience") else 0.55
        return titles, confidence

    fallback = [line for line in full_text.splitlines()[:8] if line.strip()]
    for line in fallback:
        normalized = _normalize_line(line)
        if any(keyword in normalized.lower() for keyword in TITLE_KEYWORDS):
            return [normalized], 0.4
    return [], 0.25


def _extract_industries(sections: dict[str, str], full_text: str) -> tuple[list[str], float]:
    from_section = _section_to_items(sections.get("industries", ""), max_items=15)
    if from_section:
        return from_section, 0.8

    lower_text = full_text.lower()
    inferred: list[str] = []
    for industry, keywords in INDUSTRY_KEYWORDS.items():
        if any(keyword in lower_text for keyword in keywords):
            inferred.append(industry)
    inferred = _dedupe_keep_order(inferred)
    if inferred:
        return inferred, 0.5
    return [], 0.3


def _extract_summary(
    sections: dict[str, str],
    titles: list[str],
    full_text: str,
) -> tuple[str | None, float]:
    summary = sections.get("summary", "")
    if summary:
        return summary[:1200], 0.85

    preamble_lines = [
        _normalize_line(line)
        for line in full_text.splitlines()
        if line.strip() and not _looks_like_contact_line(line)
    ]
    if preamble_lines:
        candidate = " ".join(preamble_lines[:4]).strip()
        if candidate:
            return candidate[:1200], 0.55

    if titles:
        return titles[0], 0.35
    return None, 0.2


def _extract_headline(
    sections: dict[str, str],
    titles: list[str],
    summary: str | None,
) -> tuple[str | None, float]:
    preamble = sections.get("preamble", "")
    for line in preamble.splitlines():
        candidate = _normalize_line(line)
        if not candidate or _looks_like_contact_line(candidate):
            continue
        if 2 <= len(candidate) <= 120:
            return candidate, 0.7

    if titles:
        return titles[0], 0.6
    if summary:
        return summary[:120], 0.35
    return None, 0.2


def _extract_structured_resume(text: str) -> StructuredResume:
    sections = _split_sections(text)
    skills, skills_conf = _extract_skills(sections, text)
    titles, titles_conf = _extract_titles(sections, text)
    industries, industries_conf = _extract_industries(sections, text)
    summary, summary_conf = _extract_summary(sections, titles, text)
    headline, headline_conf = _extract_headline(sections, titles, summary)

    education_items = _section_to_items(sections.get("education", ""), max_items=25)
    experience_items = _section_to_items(sections.get("experience", ""), max_items=40)

    education_json = {"items": [{"raw": item} for item in education_items]}
    experience_json = {"items": [{"raw": item} for item in experience_items]}

    confidence_json = {
        "headline": headline_conf,
        "summary": summary_conf,
        "skills": skills_conf,
        "titles": titles_conf,
        "industries": industries_conf,
        "education_json": 0.75 if education_items else 0.35,
        "experience_json": 0.8 if experience_items else 0.35,
    }
    parse_confidence = round(sum(confidence_json.values()) / len(confidence_json), 4)

    keywords = _dedupe_keep_order(skills + titles + industries, key_norm=str.lower)[:50]
    return StructuredResume(
        headline=headline,
        summary=summary,
        skills=skills,
        titles=titles,
        industries=industries,
        education_json=education_json,
        experience_json=experience_json,
        keywords=keywords,
        confidence_json=confidence_json,
        parse_confidence=parse_confidence,
    )


def _mark_failed(db: Session, resume: ResumeVersion, code: str, message: str) -> None:
    resume.status = ResumeVersionStatus.FAILED
    resume.error_code = code
    resume.error_message = message[:2000]
    resume.parsed_at = None
    resume.parse_confidence = None
    db.add(resume)
    db.commit()


def _save_extracted_text(resume: ResumeVersion, extracted_text: str) -> str | None:
    storage = get_storage()
    key = f"extracted/{resume.user_id}/{resume.id}.txt"
    payload = io.BytesIO(extracted_text.encode("utf-8"))
    try:
        return storage.put_file(key, payload)
    except Exception as exc:
        logger.warning(
            "failed to persist extracted text for resume_version_id=%s: %s",
            resume.id,
            exc,
        )
        return None


@celery_app.task(name="parse_resume")
def parse_resume(resume_version_id: str) -> dict:
    db: Session = SessionLocal()
    try:
        parsed_resume_id = uuid.UUID(resume_version_id)
    except ValueError:
        return {
            "resume_version_id": resume_version_id,
            "status": ResumeVersionStatus.FAILED.value,
            "error_code": "INVALID_RESUME_VERSION_ID",
        }

    resume: ResumeVersion | None = None
    try:
        resume = db.get(ResumeVersion, parsed_resume_id)
        if not resume:
            return {"resume_version_id": resume_version_id, "status": "missing"}

        user = db.get(User, resume.user_id)
        if user is None:
            _mark_failed(db, resume, "USER_NOT_FOUND", "owner user was not found")
            return {
                "resume_version_id": str(resume.id),
                "status": resume.status.value,
                "error_code": "USER_NOT_FOUND",
            }
        if user.status != UserStatus.ACTIVE:
            _mark_failed(db, resume, "USER_NOT_ACTIVE", "owner user is not active")
            return {
                "resume_version_id": str(resume.id),
                "status": resume.status.value,
                "error_code": "USER_NOT_ACTIVE",
            }

        resume.status = ResumeVersionStatus.SCANNING
        resume.error_code = None
        resume.error_message = None
        resume.parsed_at = None
        db.add(resume)
        db.commit()
        db.refresh(resume)

        # Stage B: malware scan stub (ClamAV integration can replace this check).
        scan_ok = True
        if not scan_ok:
            _mark_failed(db, resume, "MALWARE_DETECTED", "malware detected in uploaded file")
            return {
                "resume_version_id": str(resume.id),
                "status": resume.status.value,
                "error_code": "MALWARE_DETECTED",
            }

        resume.status = ResumeVersionStatus.PARSING
        db.add(resume)
        db.commit()
        db.refresh(resume)

        extracted_text = _extract_resume_text(resume)
        structured = _extract_structured_resume(extracted_text)

        extracted_text_uri = _save_extracted_text(resume, extracted_text)
        if extracted_text_uri:
            resume.extracted_text_uri = extracted_text_uri

        profile = db.get(Profile, resume.user_id)
        if profile is None:
            profile = Profile(user_id=resume.user_id)
        manual_fields = _manual_override_fields(profile)

        if "headline" not in manual_fields:
            profile.headline = structured.headline
        if "summary" not in manual_fields:
            profile.summary = structured.summary
        if "skills" not in manual_fields:
            profile.skills = structured.skills
        if "titles" not in manual_fields:
            profile.titles = structured.titles
        if "industries" not in manual_fields:
            profile.industries = structured.industries
        profile.education_json = structured.education_json
        profile.experience_json = structured.experience_json
        profile.keywords = structured.keywords
        profile.confidence_json = _merge_confidence_json(
            structured.confidence_json,
            profile.confidence_json,
            manual_fields,
        )
        profile.source_resume_id = resume.id
        db.add(profile)

        resume.status = ResumeVersionStatus.PARSED
        resume.parse_confidence = structured.parse_confidence
        resume.parsed_at = datetime.now(timezone.utc)
        resume.error_code = None
        resume.error_message = None
        db.add(resume)
        db.commit()

        logger.info("parse_resume completed resume_version_id=%s", resume.id)
        return {"resume_version_id": str(resume.id), "status": resume.status.value}
    except ResumeTaskError as exc:
        db.rollback()
        if resume:
            _mark_failed(db, resume, exc.code, exc.message)
            return {
                "resume_version_id": str(resume.id),
                "status": resume.status.value,
                "error_code": exc.code,
            }
        return {
            "resume_version_id": resume_version_id,
            "status": ResumeVersionStatus.FAILED.value,
            "error_code": exc.code,
        }
    except Exception as exc:
        db.rollback()
        try:
            if resume is None:
                resume = db.get(ResumeVersion, parsed_resume_id)
            if resume:
                _mark_failed(db, resume, "PARSING_ERROR", str(exc))
        except Exception:
            pass
        raise
    finally:
        db.close()
